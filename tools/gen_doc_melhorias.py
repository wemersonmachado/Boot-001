# -*- coding: utf-8 -*-
"""Gera o documento de mudanças NÃO enviadas ao Railway (melhorias do modo autônomo)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x1E, 0x7A, 0x33)
GRAY = RGBColor(0x60, 0x60, 0x60)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def h(txt, size=15, color=BLUE, before=10, after=4):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True
    r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    return p

def para(txt, color=None, bold=False, italic=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(txt, color=None, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix); rb.bold = True
        if color: rb.font.color.rgb = color
    r = p.add_run(txt)
    return p

# ── Capa ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MUDANÇAS NÃO ENVIADAS AO RAILWAY"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Trader Bot 001 — Melhorias do modo AUTÔNOMO"); r.font.size = Pt(13); r.font.color.rgb = GRAY
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("Gerado em 22/06/2026 · APENAS na pasta local · aguardando deploy"); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RED

para("")
para("⚠️ IMPORTANTE: estas alterações estão SOMENTE na pasta local "
     "(trader_001/config.py, trader_001/main.py, trader_001/database.py e "
     "trader_001/signal_filters.py). NÃO foram enviadas ao Railway. O bot em produção ainda "
     "roda a versão anterior (commit 0eb781b). Use este documento para revisar e enviar quando "
     "decidir.", color=RED, bold=True)

para("Atualização 22/06 (3ª passada): pacote de ACURÁCIA do canal SINAIS — ver Seção E. "
     "MTF e tag estrutural viraram gate, confluência mínima, Brain a partir de 55, rastreio de "
     "win-rate dos sinais e auto-tune do SINAIS. Pump/Dump NÃO foi tocado (a pedido). Testado "
     "localmente: dashboard 200 e '[SINAIS] AGGRESSIVE | 2 enviados'.", italic=True, color=GREEN)

para("Atualização 22/06 (2ª passada): janela PAPER REMOVIDA a pedido; implementados "
     "banco persistente (volume), tag de modo, loop signal_outcomes, auto-tune do score e "
     "alavancagem por volatilidade. EXCLUÍDOS a pedido: taxas/funding no PnL e backtest do "
     "motor real.", italic=True, color=BLUE)

# ── Seção A ──────────────────────────────────────────────────────────────────
h("A) Implementado agora (local) — pronto para revisão")

para("Tudo controlado por novas constantes em config.py (defaults conservadores). "
     "Compila sem erros (py_compile OK).", italic=True, color=GRAY)

h("1. Janela obrigatória em PAPER (warm-up) — REMOVIDA", size=12, color=RED, before=8, after=2)
bullet("REMOVIDA a pedido. O modo AUTÔNOMO agora opera assim que é acionado e encontra "
       "entradas conforme o perfil, sem espera inicial. Constante AUTO_PAPER_WARMUP_MIN=0.",
       bold_prefix="Status: ")

h("2. Circuit breaker da Binance (autorização no Telegram)", size=12, color=GREEN, before=8, after=2)
bullet("Após 3 ERROS consecutivos da Binance, o bot manda mensagem no Telegram pedindo "
       "autorização: /continuar (segue) ou /pausar (para agora).", bold_prefix="O quê: ")
bullet("Se NÃO responder em até 5 minutos, o bot PAUSA TUDO automaticamente (BOT_PAUSED=True).", bold_prefix="Timeout: ")
bullet("CIRCUIT_BREAKER_ENABLED=True · CB_ERROR_THRESHOLD=3 · CB_AUTH_TIMEOUT_S=300.", bold_prefix="Config: ")
bullet("_register_binance_error()/_register_binance_ok() (hook no _job_sync_binance), "
       "job_circuit_breaker_watch (scheduler 30s), comandos /continuar e /pausar.", bold_prefix="Onde: ")
bullet("Hoje o contador de erros é alimentado pelo job de sync da Binance. Pode-se estender "
       "para mais pontos (saldo, execução de ordem) numa próxima passada.", bold_prefix="Obs: ")

h("3. Anti-overtrading no mesmo ativo", size=12, color=GREEN, before=8, after=2)
bullet("Após abrir uma entrada num ativo, só permite NOVA entrada no MESMO ativo depois de "
       "15 minutos E apenas se o sinal for 'claro' (score alto).", bold_prefix="O quê: ")
bullet("SAME_ASSET_COOLDOWN_MIN=15 · CLEAR_SIGNAL_MIN_SCORE=80.", bold_prefix="Config: ")
bullet("_same_asset_blocked() + gate na branch AUTÔNOMO (bloqueio 'anti-overtrading'). "
       "Registra o ts por ativo em _asset_last_entry_ts ao abrir.", bold_prefix="Onde: ")

h("4. Teto de exposição agregada", size=12, color=GREEN, before=8, after=2)
bullet("Bloqueia novas entradas se o notional total das posições abertas exceder o teto "
       "sobre a banca.", bold_prefix="O quê: ")
bullet("MAX_TOTAL_EXPOSURE_RATIO=1.5 (notional_total / banca).", bold_prefix="Config: ")
bullet("_exposure_blocked() — gate pré-loop no job_auto_trade.", bold_prefix="Onde: ")

h("5. Teto diário de entradas", size=12, color=GREEN, before=8, after=2)
bullet("Limita o nº de entradas autônomas por dia (anti-overtrading global).", bold_prefix="O quê: ")
bullet("MAX_TRADES_PER_DAY=30 (0 = sem limite).", bold_prefix="Config: ")
bullet("_trades_today_blocked() + contador _trades_today (reseta por data UTC).", bold_prefix="Onde: ")

h("6. Banco persistente (volume)", size=12, color=GREEN, before=8, after=2)
bullet("DB_PATH agora vem da env (config.py). Aponte para um VOLUME montado (ex.: "
       "/data/trader_001.db) e o histórico/kill-switch/sessão SOBREVIVEM aos deploys.", bold_prefix="O quê: ")
bullet("No Railway: criar um Volume e definir DB_PATH=/data/trader_001.db nas Variables. "
       "Postgres completo fica como evolução futura (não necessário agora).", bold_prefix="Como ativar: ")

h("7. Tag de modo nas trades (#11)", size=12, color=GREEN, before=8, after=2)
bullet("Coluna 'mode' na tabela trades (+ migração ALTER) e gravação de OPERATION_MODE "
       "em cada abertura. Agora dá pra separar performance AUTÔNOMO × SINAIS (antes era None).", bold_prefix="O quê: ")

h("8. Loop de aprendizado signal_outcomes (#10/#14)", size=12, color=GREEN, before=8, after=2)
bullet("O fechamento por SL/TP da Binance (sync) passa a chamar record_signal_outcome — "
       "fecha o loop que alimenta o ML/score adaptativo (antes só o fechamento manual fazia).", bold_prefix="O quê: ")

h("9. Auto-tune do min_score (#12)", size=12, color=GREEN, before=8, after=2)
bullet("Job a cada 15min mede a taxa de acerto dos últimos N trades e ajusta o corte de "
       "score: win-rate baixa → mais seletivo; alta → libera um pouco. Conservador e com limites.", bold_prefix="O quê: ")
bullet("AUTOTUNE_SCORE_ENABLED, AUTOTUNE_LOOKBACK=30, AUTOTUNE_MAX_TIGHTEN=8, AUTOTUNE_MAX_LOOSEN=3.", bold_prefix="Config: ")

h("10. Alavancagem por volatilidade / ATR (#3)", size=12, color=GREEN, before=8, after=2)
bullet("Reduz a alavancagem em ativos mais voláteis (ATR% acima da referência), somando-se "
       "à redução por exposição já existente.", bold_prefix="O quê: ")
bullet("LEVERAGE_BY_VOLATILITY=True, ATR_PCT_REF=1.5, LEVERAGE_VOL_FLOOR=3.", bold_prefix="Config: ")

para("")
para("Novos comandos no Telegram: /continuar · /pausar. "
     "Novo bloco de config em config.py. Novo job no scheduler (circuit_breaker, 30s).",
     bold=True, color=BLUE)

# ── Seção B: status das 15 ────────────────────────────────────────────────────
h("B) Status das 15 melhorias sugeridas")
tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, txt in enumerate(["Melhoria", "Status", "Observação"]):
    hdr[i].paragraphs[0].add_run(txt).bold = True
rows = [
 ("Janela em PAPER (warm-up)", "REMOVIDO", "A pedido: opera assim que acionado"),
 ("Circuit breaker Binance", "IMPLEMENTADO", "3 erros → Telegram → 5min → pausa"),
 ("Anti-overtrading mesmo ativo", "IMPLEMENTADO", "15min + sinal claro (score 80)"),
 ("Teto de exposição agregada", "IMPLEMENTADO", "MAX_TOTAL_EXPOSURE_RATIO=1.5"),
 ("Teto diário de entradas", "IMPLEMENTADO", "MAX_TRADES_PER_DAY=30"),
 ("Banco persistente (volume)", "IMPLEMENTADO", "DB_PATH via env → volume (Postgres = futuro)"),
 ("Tag de modo nas trades", "IMPLEMENTADO", "Coluna 'mode' + migração + gravação"),
 ("signal_outcomes + loop ML", "IMPLEMENTADO", "Registro no fechamento por SL/TP (sync)"),
 ("Auto-tune do min_score", "IMPLEMENTADO", "job 15min ajusta corte por win-rate"),
 ("Alavancagem por volatilidade (ATR)", "IMPLEMENTADO", "Reduz lev em ativos voláteis (ATR%)"),
 ("PnL realizado no fechamento", "JÁ ENVIADO", "Corrigido no commit 0eb781b (Railway)"),
 ("Notificar stop/alvo (Binance)", "JÁ ENVIADO", "Corrigido no commit 0eb781b (Railway)"),
 ("Alavancagem adaptativa p/ exposição", "JÁ EXISTIA", "LEVERAGE ADAPTIVE + anti-martingale"),
 ("Filtros BTC veto / spread", "JÁ EXISTIA", "BTC veto + max_spread_pct por perfil"),
 ("Observabilidade (porquê do sinal)", "JÁ EXISTIA", "_blk/_blocks + EXEC-RESUMO"),
 ("Taxas+funding no PnL e no R:R", "EXCLUÍDO", "Você pediu para NÃO fazer"),
 ("Backtest do motor V6 REAL", "EXCLUÍDO", "Você pediu para NÃO fazer"),
 ("SINAIS: MTF vira gate (bloqueio)", "IMPLEMENTADO", "SINAIS_MTF_HARD_GATE (Seção E1)"),
 ("SINAIS: tag estrutural em todos perfis", "IMPLEMENTADO", "SINAIS_REQUIRE_STRUCT_ALL (E2)"),
 ("SINAIS: confluência mínima de tags", "IMPLEMENTADO", "SINAIS_MIN_CONFLUENCE (E3)"),
 ("SINAIS: Brain a partir de 55", "IMPLEMENTADO", "SINAIS_BRAIN_MIN_SCORE=55 (E4)"),
 ("SINAIS: rastreio de win-rate", "IMPLEMENTADO", "signal_outcomes + watcher 3min (E5)"),
 ("SINAIS: auto-tune do corte", "IMPLEMENTADO", "job 20min por acerto medido (E6)"),
 ("Pump/Dump (peso/cooldown)", "NÃO TOCADO", "A pedido: manter como está"),
]
for nome, status, obs in rows:
    c = tbl.add_row().cells
    c[0].text = nome
    rs = c[1].paragraphs[0].add_run(status)
    rs.bold = True
    rs.font.color.rgb = GREEN if status == "IMPLEMENTADO" else (BLUE if status in ("JÁ ENVIADO", "JÁ EXISTIA") else RED)
    c[2].text = obs

# ── Seção C: excluídos + evolução futura ──────────────────────────────────────
h("C) Excluídos a pedido + evolução futura (opcional)")
para("Excluídos por sua decisão (NÃO implementados):", bold=True, color=RED)
bullet("Taxas+funding no PnL e no R:R líquido (custo de ida-e-volta na decisão de entrada).", bold_prefix="#6/#7 ")
bullet("Backtest/walk-forward do motor V6 real sobre candles históricos.", bold_prefix="#9 ")
para("")
para("Evolução futura (opcional, quando quiser):", bold=True, color=BLUE)
bullet("Postgres completo no lugar do SQLite-em-volume (mais robusto para concorrência/escala).", bold_prefix="Banco: ")
bullet("Retreino automático do ML a partir do signal_outcomes (agora que o loop está populado).", bold_prefix="ML: ")

# ── Seção D: como enviar + teste ──────────────────────────────────────────────
h("D) Como enviar ao Railway depois + checklist de teste")
para("Quando decidir enviar (não fazer agora, conforme pedido):", bold=True)
bullet("Copiar trader_001/config.py, trader_001/main.py e trader_001/database.py para Desktop/Trade/Boot-001_repo/")
bullet("cd Boot-001_repo → git add config.py main.py database.py → git commit -m \"...\" → git push origin main")
bullet("Para o banco persistir: criar um Volume no Railway e definir DB_PATH=/data/trader_001.db nas Variables.")
bullet("A Railway reconstrói sozinha (~3-6 min).")
para("")
para("Antes de operar REAL com isto:", bold=True, color=RED)
bullet("Testar com PAPER_TRADING=ON primeiro (valida gates sem risco).")
bullet("Confirmar que /continuar e /pausar respondem no Telegram.")
bullet("Conferir os defaults (warm-up 15min, exposição 1.5x, 30 trades/dia, score claro 80) "
       "e ajustar ao seu gosto antes do deploy.")
bullet("Lembrar: o DB da Railway zera no deploy — o teto diário e a sessão recomeçam do zero.")

# ── Seção E: Acurácia do canal SINAIS (3ª passada) ───────────────────────────
h("E) Acurácia do canal SINAIS (2026-06-22, 3ª passada)")
para("Objetivo: aumentar a ASSERTIVIDADE dos sinais transmitidos no modo SINAIS, em "
     "TODOS os perfis, SEM mexer no Pump/Dump (a pedido) e SEM tocar no autônomo/real. "
     "evaluate_signal() é chamado SOMENTE pelo job_sinais_scan, então todos os gates abaixo "
     "afetam apenas o canal de sinais.", italic=True, color=GRAY)
para("Arquivos: trader_001/config.py, trader_001/signal_filters.py, trader_001/database.py, "
     "trader_001/main.py.", italic=True, color=GRAY)

h("E1. MTF vira GATE (bloqueio), não só penalidade", size=12, color=GREEN, before=8, after=2)
bullet("Se o timeframe superior DIVERGE da direção do sinal, o sinal é BLOQUEADO. Antes "
       "levava só -5/-8 pts e ainda passava se o score fosse alto. Ausência de TF superior "
       "no cache permanece neutra.", bold_prefix="O quê: ")
bullet("SINAIS_MTF_HARD_GATE=True. Gate em signal_filters.evaluate_signal (após check_mtf).", bold_prefix="Config/Onde: ")

h("E2. Tag estrutural V6 obrigatória em TODOS os perfis", size=12, color=GREEN, before=8, after=2)
bullet("Antes só o NORMAL exigia tag estrutural; o AGGRESSIVE (padrão do canal) deixava "
       "passar momentum 'pelado'. Agora Conservador e Agressivo também exigem ≥1 tag "
       "(BOS/OB-FVG/sweep/FIB/divergência/...).", bold_prefix="O quê: ")
bullet("SINAIS_REQUIRE_STRUCT_ALL=True.", bold_prefix="Config: ")

h("E3. Confluência mínima (nº de tags distintas)", size=12, color=GREEN, before=8, after=2)
bullet("Bloqueia o 'score alto solitário' (1 fator inflado): exige N tags estruturais "
       "DISTINTAS concordando. count_structural_tags() conta da tag mais longa p/ a mais "
       "curta (evita FIB618 recontar como FIB).", bold_prefix="O quê: ")
bullet("SINAIS_MIN_CONFLUENCE = {CONSERVATIVE:2, NORMAL:2, AGGRESSIVE:1}.", bold_prefix="Config: ")

h("E4. Claude Brain a partir do score 55 (era 65)", size=12, color=GREEN, before=8, after=2)
bullet("Sinais de 55-64 do Agressivo escapavam da IA. Agora o Brain (quando ligado) avalia "
       "a partir de 55. Obs: Brain vem DESLIGADO por padrão — sem custo/latência até ativar.", bold_prefix="O quê: ")
bullet("SINAIS_BRAIN_MIN_SCORE=55.", bold_prefix="Config: ")

h("E5. Rastreio de resultado dos sinais (win-rate medível)", size=12, color=GREEN, before=8, after=2)
bullet("Cada sinal transmitido é registrado (entry/SL/TP1/score/tags). Um job a cada 3min "
       "(job_sinais_outcome_watch) compara os candles após a entrada: WIN se tocou o alvo "
       "primeiro, LOSS se tocou o stop, TIMEOUT após 6h. Grava em signal_outcomes. Antes só "
       "TRADES executados eram medidos — sinais nunca tinham win-rate.", bold_prefix="O quê: ")
bullet("SINAIS_OUTCOME_TRACKING=True, SINAIS_OUTCOME_MAX_AGE_H=6. Conservador: em empate "
       "TP+SL no mesmo candle conta LOSS. Estado em memória (não persiste reinício).", bold_prefix="Config/Obs: ")

h("E6. Auto-tune do corte do SINAIS por acerto medido", size=12, color=GREEN, before=8, after=2)
bullet("job_sinais_autotune (a cada 20min) lê o win-rate dos últimos N sinais resolvidos e "
       "ajusta _sinais_score_offset: win-rate baixa → corte mais alto; alta → libera. "
       "Independente do auto-tune dos trades.", bold_prefix="O quê: ")
bullet("SINAIS_AUTOTUNE_ENABLED=True, LOOKBACK=30, MAX_TIGHTEN=8, MAX_LOOSEN=3. "
       "get_recent_signal_stats() em database.py.", bold_prefix="Config/Onde: ")

para("")
para("NÃO mexido (a pedido): Pump/Dump — orçamento, cooldown e tamanho permanecem como estão.",
     bold=True, color=RED)

h("E7. Teste local executado (22/06) — resultado", size=12, color=BLUE, before=8, after=2)
bullet("Bot ligado localmente (DB temporária, VIP/CANAL blindados → sinais só ao chat pessoal), "
       "modo SINAIS perfil AGRESSIVO ativado pelo dashboard.", bold_prefix="Setup: ")
bullet("Dashboard / e endpoints (/settings, /signals/latest, /trades/active, /claude-brain/status, "
       "/alerts/pump_dump) responderam 200. Startup completo, SEM tracebacks.", bold_prefix="Dashboard: ")
bullet("Pipeline rodou ponta-a-ponta: '[SINAIS] AGGRESSIVE | 2 enviados (pd=0 reg=2)' — 2 sinais "
       "passaram os novos gates (de 8-11 candidatos): seletividade coerente, sem zerar o fluxo.", bold_prefix="SINAIS: ")
bullet("/performance respondeu 200 porém LENTO (~22s) e houve congestionamento de scheduler no "
       "cold-start (chamadas síncronas à Binance no event-loop). PRÉ-EXISTENTE, não introduzido "
       "por estas mudanças — candidato a otimização futura (asyncio.to_thread nos pontos síncronos).", bold_prefix="Achado: ")

h("E8. Como enviar ao Railway (junto com o resto)", size=12, color=BLUE, before=8, after=2)
bullet("Copiar TAMBÉM trader_001/signal_filters.py (além de config.py, main.py, database.py) "
       "para Boot-001_repo/ e commitar/empurrar.")
bullet("Todos os gates são controlados por flags em config.py (default ligado) — para reverter "
       "qualquer um, basta setar a flag correspondente para False.")

# ── Seção F: Circuit breaker por PERDAS (local, pós-deploy 057ab98) ───────────
h("F) Circuit breaker por TRADES PERDEDORES (2026-06-22) — DEPLOYADO (aaade49)")
para("NOTA: A–E foram ao Railway no commit 057ab98; esta Seção F foi ao Railway no commit "
     "aaade49 (gatilho por perdas + DESLIGAR o gatilho por erros da Binance, que mandava a "
     "mensagem '3 erros seguidos' constante).", italic=True, color=GREEN)
para("Correção de intenção: o circuit breaker fora pedido para '3 erros', mas o que você "
     "queria era 3 TRADES PERDEDORES seguidos (proteção de banca), não 3 erros de API.",
     bold=True, color=RED)
bullet("Novo GATILHO PRINCIPAL: após CB_LOSS_THRESHOLD (=3) trades perdedores seguidos, "
       "dispara o circuit breaker — mesma mensagem no Telegram (/continuar ou /pausar; sem "
       "resposta em 5min -> pausa tudo).", bold_prefix="O quê: ")
bullet("O gatilho por ERROS da Binance (CB_ERROR_THRESHOLD) foi MANTIDO como secundário "
       "(falha técnica), agora com mensagem própria. Ambos usam o mesmo fluxo _cb_arm().", bold_prefix="Mantido: ")
bullet("Usa _consecutive_losses (o mesmo do anti-martingale). _cb_loss_ack_streak evita "
       "re-alertar a cada nova perda após /continuar — só re-arma a cada novo bloco de 3 "
       "perdas. Reseta sozinho quando o streak é zerado por vitórias.", bold_prefix="Como: ")
bullet("config.py: CB_LOSS_THRESHOLD=3. main.py: _cb_arm(), _maybe_trip_loss_breaker(), "
       "chamada após cada incremento de _consecutive_losses no fechamento (DCA e normal), "
       "_cb_resume() reconhece o baseline.", bold_prefix="Onde: ")
bullet("Testado local: 1-2 perdas não arma; 3a arma ('3 trades perdedores seguidos'); "
       "após /continuar, 4a não re-arma e a 6a re-arma. py_compile OK.", bold_prefix="Teste: ")
bullet("Enviar junto: trader_001/config.py e trader_001/main.py -> railway_repo/ -> push.", bold_prefix="Deploy: ")

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "Mudancas_Nao_Enviadas_Railway.docx")
doc.save(out)
print("Documento gerado:", out)
