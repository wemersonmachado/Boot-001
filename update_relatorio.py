import os
from docx import Document

filepath = r"C:\Users\welli\Desktop\Trade\Trader_Bot_001\PDF\RELATÓRIO DE MELHORIAS.docx"

try:
    if os.path.exists(filepath):
        doc = Document(filepath)
    else:
        doc = Document()
        p = doc.add_paragraph()
        p.add_run('RELATÓRIO DE MELHORIAS').bold = True
        p.add_run('\n\n')

    p2 = doc.add_paragraph()
    p2.add_run('Atualização - 20 de Junho de 2026').bold = True
    
    p = doc.add_paragraph()
    p.add_run('1. Motor Principal (main.py) - Correção de Bug do PnL\n').bold = True
    doc.add_paragraph('Problema: Trades fechados estavam registrando 0.0 no PnL do banco de dados, cegando as métricas de acompanhamento.')
    doc.add_paragraph('Solução: Foi injetada uma lógica de recálculo exato de PnL usando o preço de saída real da exchange momentos antes de escrever no banco de dados.')
    
    p = doc.add_paragraph()
    p.add_run('2. Motor de Sinais (signal_filters.py) - Hard Block de RSI Extremo\n').bold = True
    doc.add_paragraph('Problema: O bot penalizava, mas ainda permitia a abertura de posições contra tendências extremas, como shorts no fundo ou longs no topo.')
    doc.add_paragraph('Solução: Substituímos a penalidade de pontos por um filtro inflexível: Sinais de LONG são bloqueados e cancelados se RSI for maior que 60. Sinais de SHORT são bloqueados se RSI for menor que 40.')

    doc.add_paragraph('-'*50)
    doc.save(filepath)
    print("Documento salvo com sucesso em:", filepath)
except Exception as e:
    print("Erro ao salvar:", e)
